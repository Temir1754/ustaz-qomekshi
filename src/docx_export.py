"""Генерация .docx КТП по структуре реального школьного шаблона
(разобран из примера: 8 КТЖ алг+геом 2024-25.docx).

Формат таблицы: № | Бөлім | р/с | Сабақтардың тақырыбы | Оқудың мақсаттары |
Сағат саны | Мерзімі | Ескерту, с группировкой по четвертям. Шапка (класс/предмет/
часы) — объединённая первая строка самой таблицы, как в реальном образце.

Темы уроков и коды целей обучения в этой версии — ЗАГЛУШКИ (нет LLM и нет
реальной программы предмета для опоры), это осознанно: подставлять
выдуманные официальные коды целей ("8.1.1.1" и т.п.) без реальной сверки
с типовой программой — то же самое рискованное поведение, которого мы
избегаем для оценок. Здесь важна корректная СТРУКТУРА документа, готовая
к тому, чтобы контент заполнила LLM с опорой на реальную программу.
"""

from __future__ import annotations

import datetime
import io

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from ktp_notes import get_explanatory_note

# Доля часов по четвертям — взята из пропорций реального образца
# (24/102 ≈ 0.235 в 1-й четверти и т.д.), с округлением остатка в последнюю четверть.
QUARTER_RATIOS = [0.24, 0.24, 0.29, 0.23]

# Весь текст структуры документа (заголовки таблицы, подписи полей, названия
# этапов КСП) должен быть на том же языке, что и содержание, которое для
# темы/этапов сгенерировал AI (см. ai_generate.py, параметр lang) — иначе
# в документе казахские подписи шаблона перемешиваются с русским содержанием
# или наоборот (баг, который был в предыдущей версии: шаблон был жёстко
# закодирован на казахском независимо от языка обучения). lang здесь — это
# язык самого документа, а не язык интерфейса бота.
DOCX_STRINGS = {
    "kk": {
        "ktp_title": (
            "{cls} сынып «{subject_label}» пәні бойынша "
            "күнтізбелік-тақырыптық жоспар {school_year} оқу жылы"
        ),
        "ktp_meta": "Аптасына {hours_per_week} сағат       Оқу жылында {total_hours} сағат",
        "ktp_textbook": "Оқулық: {textbook}",
        "ktp_headers": ["№", "Бөлім", "р/с", "Сабақтардың тақырыбы", "Оқудың мақсаттары", "Сағат саны", "Мерзімі", "Ескерту"],
        "ktp_quarter": "{q_index}-ТОҚСАН ({q_hours} сағат)",
        "ktp_stub_lesson": "Тема {n} — толтыру қажет (нақты тақырыпты AI/мұғалім енгізеді)",
        "ksp_title": "Қысқа мерзімді сабақ жоспары",
        "ksp_info_labels": [
            "Бөлім:",
            "Педагогтің аты-жөні:",
            "Күні:",
            "Сынып:",
            "Сабақтың тақырыбы:",
            "Оқу бағдарламасына сәйкес оқыту мақсаттары:",
            "Сабақтың мақсаты:",
        ],
        "ksp_goal_default": (
            "Оқушылар «{topic}» тақырыбын меңгереді және тиісті есептерді шеше алады "
            "(толық тұжырымдау қажет)."
        ),
        "ksp_objectives_placeholder": "— (нақты бағдарламадан алынуы қажет)",
        "ksp_headers": ["Уақыты/кезеңдері", "Педагогтің әрекеті", "Оқушының әрекеті", "Бағалау", "Ресурстар"],
        "ksp_stages": [
            (
                "Ұйымдастыру кезеңі\n(5 мин)",
                "Амандасу, сабаққа психологиялық дайындық жасау, үй тапсырмасын тексеру.",
                "Мұғаліммен амандасады, үй жұмысын тексереді, сұрақтарға жауап береді.",
                "Ауызша мақтау («Өте жақсы!», «Дұрыс!»)",
                "Тақта, слайд",
            ),
            (
                "Сабақтың басы\n(жеке жұмыс)",
                "«{topic}» тақырыбы бойынша кіріспе тапсырма береді.",
                "Тапсырманы дербес орындайды, қиындық туса мұғалімнен сұрайды.",
                "Дескриптор бойынша — 1 балл (толтыру қажет)",
                "Слайд",
            ),
            (
                "Сабақтың ортасы\n(жұптық/топтық жұмыс)",
                "Деңгейлеп тапсырма береді (жеңіл / орташа / күрделі).",
                "Жұппен немесе топпен талқылап шешеді, тақтада көрсетеді.",
                "Дескриптор бойынша — 2 балл (толтыру қажет)",
                "Жұмыс парағы",
            ),
            (
                "Сабақтың соңы\n(5 мин)",
                "Кері байланыс алады, үй тапсырмасын береді.",
                "Түсінгенін/түсінбегенін білдіреді («Түсіндім» / «Сұрағым бар»).",
                "Жинаған балл бойынша бағалау",
                "—",
            ),
        ],
    },
    "ru": {
        "ktp_title": (
            "Календарно-тематическое планирование по предмету «{subject_label}», "
            "{cls} класс, {school_year} учебный год"
        ),
        "ktp_meta": "В неделю: {hours_per_week} ч.       За учебный год: {total_hours} ч.",
        "ktp_textbook": "Учебник: {textbook}",
        "ktp_headers": ["№", "Раздел", "№ урока", "Тема урока", "Цели обучения", "Кол-во часов", "Дата", "Примечание"],
        "ktp_quarter": "{q_index} ЧЕТВЕРТЬ ({q_hours} ч.)",
        "ktp_stub_lesson": "Тема {n} — требуется заполнить (тему добавит AI/учитель)",
        "ksp_title": "Краткосрочный план урока",
        "ksp_info_labels": [
            "Раздел:",
            "ФИО педагога:",
            "Дата:",
            "Класс:",
            "Тема урока:",
            "Цели обучения согласно учебной программе:",
            "Цель урока:",
        ],
        "ksp_goal_default": (
            "Учащиеся освоят тему «{topic}» и научатся решать соответствующие задачи "
            "(нужна более точная формулировка)."
        ),
        "ksp_objectives_placeholder": "— (нужно взять из реальной программы)",
        "ksp_headers": ["Этап урока", "Действия педагога", "Действия ученика", "Оценивание", "Ресурсы"],
        "ksp_stages": [
            (
                "Организационный этап\n(5 мин)",
                "Приветствие, психологический настрой на урок, проверка домашнего задания.",
                "Приветствует учителя, показывает домашнее задание, отвечает на вопросы.",
                "Устная похвала («Отлично!», «Верно!»)",
                "Доска, слайд",
            ),
            (
                "Начало урока\n(индивидуальная работа)",
                "Даёт вводное задание по теме «{topic}».",
                "Самостоятельно выполняет задание, при затруднении обращается к учителю.",
                "По дескриптору — 1 балл (нужно заполнить)",
                "Слайд",
            ),
            (
                "Середина урока\n(парная/групповая работа)",
                "Даёт задание с уровневой дифференциацией (лёгкий / средний / сложный).",
                "Обсуждает и решает в паре или группе, представляет на доске.",
                "По дескриптору — 2 балла (нужно заполнить)",
                "Рабочий лист",
            ),
            (
                "Конец урока\n(5 мин)",
                "Собирает обратную связь, даёт домашнее задание.",
                "Сообщает, понял ли материал («Понятно» / «Есть вопрос»).",
                "Оценивание по набранным баллам",
                "—",
            ),
        ],
    },
}


def _s(lang: str) -> dict:
    return DOCX_STRINGS.get(lang, DOCX_STRINGS["kk"])


def _landscape(doc: Document) -> None:
    """Оба реальных образца (КТП и КСП) — альбомной ориентации: широкие таблицы
    (8 колонок у КТП, 5 у КСП) иначе не помещаются по ширине."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def _set_base_font(doc: Document, name: str = "Times New Roman", size: int = 12) -> None:
    """Шрифт всего документа по умолчанию — то, что явно не переопределено
    другим run.font (как заголовки таблицы, см. _header_cell)."""
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    # без этого Word показывает кириллицу шрифтом по умолчанию, даже если
    # w:rFonts/w:ascii уже стоит на Times New Roman.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _bold(cell) -> None:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True


def _header_cell(cell, text: str, size: int = 11) -> None:
    """Заголовок таблицы (шапка колонок) — Times New Roman, жирный, size пт
    (по умолчанию 11), без заливки — отдельно от общего шрифта документа (12 пт)."""
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)


def _shade(cell, hex_color: str = "D9E2F3") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


_WEEKDAY_PATTERNS = {
    1: [0],
    2: [0, 3],
    3: [0, 2, 4],
    4: [0, 1, 3, 4],
    5: [0, 1, 2, 3, 4],
}


def lesson_weekdays(hours_per_week: int) -> list[int]:
    """Дефолтный набор из hours_per_week будних дней (0=Пн..4=Пт), равномерно
    распределённых по неделе — используется как предзаполненный выбор в UI
    бота (bot.py), который учитель может изменить кнопками; здесь — только
    запасной вариант, если вызывающий код не передал weekdays явно."""
    if hours_per_week in _WEEKDAY_PATTERNS:
        return _WEEKDAY_PATTERNS[hours_per_week]
    if hours_per_week < 1:
        return [0]
    return [0, 1, 2, 3, 4]


def _generate_lesson_dates(total_hours: int, start_date: datetime.date, weekdays: list[int]) -> list[datetime.date]:
    """total_hours дат подряд начиная с start_date, только по дням недели из
    weekdays — без учёта каникул между четвертями (реального календаря школы
    у нас нет), это приближение по словам учителя (start_date/weekdays), а не
    точное расписание из Kundelik.kz."""
    weekday_set = set(weekdays) or {0}
    d = start_date
    dates: list[datetime.date] = []
    while len(dates) < total_hours:
        if d.weekday() in weekday_set:
            dates.append(d)
        d += datetime.timedelta(days=1)
    return dates


def build_ktp_docx(
    cls: str,
    subject_label: str,
    hours_per_week: int,
    total_hours: int,
    school_year: str = "2025-2026",
    outline: list[tuple[str, str]] | None = None,
    lang: str = "kk",
    textbook: str | None = None,
    start_date: datetime.date | None = None,
    weekdays: list[int] | None = None,
) -> bytes:
    """outline: список (раздел, тема_урока) длиной total_hours — из ai_generate.generate_ktp_outline.
    Если не передан, темы заполняются заглушками (без LLM). lang — язык самого документа
    (должен совпадать с языком, на котором generate_ktp_outline сгенерировал темы). textbook —
    необязательная строка «учебник/издательство», которую вводит сам учитель (не выдумывается AI).
    start_date/weekdays — с какой даты и по каким дням недели (0=Пн..4=Пт) расставлять даты
    уроков (см. bot.py: учитель указывает в тексте или подтверждает кнопками). Если не переданы —
    по умолчанию 1 сентября school_year и lesson_weekdays(hours_per_week).
    Если для (subject_label, lang) есть сохранённый официальный текст «Түсінік хат»
    (см. ktp_notes.py), он вставляется перед таблицей — так же, как в реальном образце."""
    strings = _s(lang)
    doc = Document()
    _landscape(doc)
    _set_base_font(doc)

    note_paragraphs = get_explanatory_note(subject_label, lang)
    if note_paragraphs:
        for i, text in enumerate(note_paragraphs):
            p = doc.add_paragraph()
            run = p.add_run(text)
            if i == 0:
                run.bold = True
                run.font.size = Pt(13)
        doc.add_paragraph()

    # Шапка (класс/предмет/часы) — не отдельные абзацы над таблицей, а объединённая
    # первая строка самой таблицы: так устроено в реальном школьном образце
    # (8 КТЖ алг+геом 2024-25.docx), а не отдельным текстом сверху.
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"

    banner = table.rows[0].cells[0].merge(table.rows[0].cells[7])
    banner_p = banner.paragraphs[0]
    banner_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = banner_p.add_run(
        strings["ktp_title"].format(cls=cls, subject_label=subject_label, school_year=school_year)
    )
    title_run.bold = True
    title_run.font.size = Pt(13)
    banner_p.add_run().add_break()
    meta_run = banner_p.add_run(strings["ktp_meta"].format(hours_per_week=hours_per_week, total_hours=total_hours))
    meta_run.italic = True
    if textbook:
        banner_p.add_run().add_break()
        textbook_run = banner_p.add_run(strings["ktp_textbook"].format(textbook=textbook))
        textbook_run.italic = True

    headers = strings["ktp_headers"]
    hdr_cells = table.rows[1].cells
    for idx, text in enumerate(headers):
        _header_cell(hdr_cells[idx], text)

    if outline is None:
        outline = [("", strings["ktp_stub_lesson"].format(n=i + 1)) for i in range(total_hours)]

    if start_date is None:
        try:
            start_date = datetime.date(int(school_year.split("-")[0]), 9, 1)
        except (ValueError, IndexError):
            start_date = datetime.date(datetime.date.today().year, 9, 1)
    lesson_dates = _generate_lesson_dates(total_hours, start_date, weekdays or lesson_weekdays(hours_per_week))
    lesson_no = 0
    hours_left = total_hours
    idx = 0
    for q_index, ratio in enumerate(QUARTER_RATIOS, start=1):
        is_last = q_index == len(QUARTER_RATIOS)
        q_hours = hours_left if is_last else round(total_hours * ratio)
        hours_left -= q_hours
        if q_hours <= 0:
            continue

        q_row = table.add_row().cells
        # № (первая колонка) остаётся отдельной ячейкой — как в реальном образце.
        merged = q_row[1].merge(q_row[7])
        merged.text = strings["ktp_quarter"].format(q_index=q_index, q_hours=q_hours)
        _bold(merged)
        _shade(merged, "F2F2F2")

        for _ in range(q_hours):
            if idx >= len(outline):
                break
            section, lesson_name = outline[idx]
            idx += 1
            lesson_no += 1
            row = table.add_row().cells
            row[0].text = ""
            row[1].text = section
            row[2].text = str(lesson_no)
            row[3].text = lesson_name
            row[4].text = "—"
            row[5].text = "1"
            row[6].text = lesson_dates[lesson_no - 1].strftime("%d.%m.%y") if lesson_no - 1 < len(lesson_dates) else ""
            row[7].text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- КСП / ҚМЖ (краткосрочный план урока) ----------
# Структура разобрана из реального образца: "Жанаманың теңдеуі 16.03.docx"
# (шапка с реквизитами + таблица из 4 этапов урока: организация / начало /
# середина / конец, с колонками "действие учителя / действие ученика /
# оценивание / ресурсы"). Код цели обучения оставлен пустым осознанно —
# это официальный код из учебной программы, выдумывать его нельзя.


def resolve_ksp_content(topic: str, ai_content: dict | None, lang: str = "kk") -> dict:
    """Сливает ai_content с дефолтными формулировками этапов на языке lang
    (DOCX_STRINGS[lang]["ksp_stages"]) в единую структуру {"goal": str, "stages":
    [{"stage_name","teacher_action","student_action","assessment","resources"},
    ... ровно 4]} — общий источник и для .docx, и для просмотра/редактирования
    в боте (bot.py), чтобы оба места не расходились и не смешивали языки."""
    strings = _s(lang)
    ai_content = ai_content or {}
    goal = ai_content.get("goal") or strings["ksp_goal_default"].format(topic=topic)
    ai_stages = ai_content.get("stages") or []
    stages = []
    for idx, (stage_name, default_teacher, default_student, default_assessment, default_resources) in enumerate(
        strings["ksp_stages"]
    ):
        s = ai_stages[idx] if idx < len(ai_stages) else {}
        stages.append(
            {
                "stage_name": stage_name,
                "teacher_action": s.get("teacher_action") or default_teacher.format(topic=topic),
                "student_action": s.get("student_action") or default_student,
                "assessment": s.get("assessment") or default_assessment,
                "resources": s.get("resources") or default_resources,
            }
        )
    return {"goal": goal, "stages": stages}


def build_ksp_docx(
    cls: str,
    subject_label: str,
    topic: str,
    teacher_name: str = "—",
    date_str: str | None = None,
    ai_content: dict | None = None,
    lang: str = "kk",
) -> bytes:
    """ai_content: {"goal": str, "stages": [{"teacher_action","student_action","assessment","resources"}, ...]}
    из ai_generate.generate_ksp_content, или уже слитый результат resolve_ksp_content().
    Если не передан, используются заглушки (без LLM). lang — язык самого документа
    (должен совпадать с языком, на котором сгенерирован ai_content)."""
    strings = _s(lang)
    date_str = date_str or datetime.date.today().strftime("%d.%m.%Y")
    doc = Document()
    _landscape(doc)
    _set_base_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(strings["ksp_title"])
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    content = resolve_ksp_content(topic, ai_content, lang)
    values = [subject_label, teacher_name, date_str, cls, topic, strings["ksp_objectives_placeholder"], content["goal"]]
    info_rows = list(zip(strings["ksp_info_labels"], values))
    info = doc.add_table(rows=len(info_rows), cols=2)
    info.style = "Table Grid"
    for i, (label, value) in enumerate(info_rows):
        info.rows[i].cells[0].text = label
        _bold(info.rows[i].cells[0])
        info.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        info.rows[i].cells[1].text = value

    doc.add_paragraph()

    headers = strings["ksp_headers"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    for stage in content["stages"]:
        row = table.add_row().cells
        row[0].text = stage["stage_name"]
        row[1].text = stage["teacher_action"]
        row[2].text = stage["student_action"]
        row[3].text = stage["assessment"]
        row[4].text = stage["resources"]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
